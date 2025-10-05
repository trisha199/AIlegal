from typing import Dict
from docx import Document
from src.utils.io import out_path

SECTIONS = [
    "Heading","Background and Facts","Liability","Injuries and Treatment",
    "Economic Damages","Non-Economic Damages","Demand","Citations and Exhibits"
]

def write_docx(case_id: str, content: str, out_dir: str) -> str:
    doc = Document()
    for line in content.splitlines():
        if line.strip() in SECTIONS:
            doc.add_heading(line.strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line)
    path = out_path(out_dir, f"demand_letter_{case_id}.docx")
    doc.save(path)
    return path

def generate_text_from_context(facts: Dict, passages: Dict) -> str:
    return "\n".join(SECTIONS) + "\n"
